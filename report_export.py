"""DOCX / XLSX report export for PackAssist.

Mirrors the PDF report (report-generator.js) section for section with the
SAME content and value formatting: part info, pack metrics, warnings,
box info, pallet, components, cost summary, supplier, approvals, comments.
Word/Excel reflow the layout (mm-exact layout is impossible outside the
browser print pipeline), but every value and table of the PDF is present
in the same order.
"""

import io
import base64


def _png_bytes(data_url):
    """Decode a PNG data URL (or raw base64) into bytes; None if unusable."""
    if not data_url or not isinstance(data_url, str):
        return None
    if data_url.startswith("data:"):
        data_url = data_url.split(",", 1)[1] if "," in data_url else ""
    try:
        raw = base64.b64decode(data_url)
    except Exception:
        return None
    return raw if len(raw) > 100 else None


def _fmt(v, digits=0):
    if v is None:
        return "—"
    try:
        return f"{float(v):,.{digits}f}".replace(",", " ")
    except (TypeError, ValueError):
        return str(v)


def _fmt_mm(d):
    if d is None:
        return "—"
    try:
        return f"{float(d):g}"
    except (TypeError, ValueError):
        return str(d)


def _fmt_euro(v, digits=2):
    if v is None:
        return "—"
    try:
        return f"{float(v):,.{digits}f} €".replace(",", " ")
    except (TypeError, ValueError):
        return str(v)


def _fmt_kg(v, digits=3):
    if v is None:
        return "—"
    try:
        return f"{float(v):,.{digits}f} kg".replace(",", " ")
    except (TypeError, ValueError):
        return str(v)


def _mode_label(data):
    mode = data.get("mode", "bulk")
    labels = {"bulk": "Gravetat", "gpu": "Optimitzat", "fast": "Planar"}
    label = labels.get(mode, mode)
    if mode == "gpu" and data.get("gpuMethod"):
        label += f" · {data['gpuMethod']}"
    return label


def _box_dims_text(data):
    bd = data.get("boxDims") or {}
    return (f"{_fmt_mm(bd.get('length'))} × {_fmt_mm(bd.get('width'))} × "
            f"{_fmt_mm(bd.get('height'))} mm")


def _dims_text(data):
    pd = data.get("pieceDims") or {}
    return f"{_fmt_mm(pd.get('l'))} × {_fmt_mm(pd.get('w'))} × {_fmt_mm(pd.get('h'))} mm"


def _build_payload(data):
    part = data.get("part") or {}
    supplier = data.get("supplier") or {}
    cost = data.get("cost") or {}
    pallet = data.get("pallet") or {}
    approvals = data.get("approvals") or {}
    items = data.get("items") or []
    piece_count = int(data.get("pieceCount") or 0)
    fill = data.get("fillPct")
    fill_rate = f"{_fmt(fill, 1)} %" if fill is not None else "—"
    total_weight = data.get("estimatedTotalWeight")
    interlocked = data.get("interlocked") or {}
    interlock_count = int(interlocked.get("count") or 0)
    show_interlock = data.get("showInterlockWarning", True) is not False
    trays = data.get("trays") or []

    item_rows = []
    for it in items:
        dims = "—"
        if it.get("l") is not None or it.get("w") is not None or it.get("h") is not None:
            dims = f"{_fmt_mm(it.get('l'))} × {_fmt_mm(it.get('w'))} × {_fmt_mm(it.get('h'))}"
        qty = it.get("qty")
        price = it.get("price")
        cost_v = (qty or 0) * (price or 0) if qty is not None and price is not None else None
        item_rows.append({
            "desc": it.get("desc") or "",
            "material": it.get("material") or "—",
            "dims": dims,
            "qty": _fmt(qty) if qty is not None else "—",
            "price": _fmt_euro(price) if price is not None else "—",
            "cost": _fmt_euro(cost_v) if cost_v is not None else "—",
        })

    has_items = bool(item_rows)
    box_cost = cost.get("boxCost")
    packaging_cost = cost.get("packagingCost")
    freight_cost = cost.get("freightCost")
    cost_per_part = cost.get("costPerPart")
    if cost_per_part is None and (box_cost is not None or packaging_cost is not None):
        total_costs = (box_cost or 0) + (packaging_cost or 0)
        cost_per_part = total_costs / piece_count if piece_count > 0 else None
    has_cost = any(v is not None for v in (box_cost, packaging_cost, freight_cost, cost_per_part))

    has_pallet = any(pallet.get(k) is not None for k in ("l", "w", "h", "weight", "boxes"))
    has_supplier = any(supplier.get(k) for k in ("name", "address", "contact", "phone", "email", "function"))
    has_approvals = any(approvals.get(k) for k in ("createdBy", "conceptName", "finalName"))

    return {
        "part": part, "supplier": supplier, "cost": cost, "pallet": pallet,
        "approvals": approvals, "items": item_rows, "has_items": has_items,
        "piece_count": piece_count, "fill_rate": fill_rate,
        "total_weight": total_weight, "mode_label": _mode_label(data),
        "box_dims": _box_dims_text(data), "dims": _dims_text(data),
        "piece_weight": data.get("pieceWeight"), "max_weight": data.get("maxWeight"),
        "interlock_count": interlock_count, "show_interlock": show_interlock,
        "tray_count": len(trays),
        "tray_pieces": [t.get("pieces") for t in trays],
        "has_pallet": has_pallet, "has_supplier": has_supplier,
        "has_cost": has_cost, "has_approvals": has_approvals,
        "box_cost": box_cost, "packaging_cost": packaging_cost,
        "freight_cost": freight_cost, "cost_per_part": cost_per_part,
        "name": data.get("stlFileName") or "PackAssist",
        "comments": data.get("comments") or "",
    }


# ─────────────────────────── DOCX ───────────────────────────

def _add_views_docx(doc, views):
    """Embed the PDF's 3D renders: hero (full width) + 3 small views in a
    row, mirroring the PDF page-1 layout."""
    if not views:
        return
    from docx.shared import Cm

    hero = _png_bytes(views.get("hero"))
    if hero:
        try:
            doc.add_picture(io.BytesIO(hero), width=Cm(16.5))
            doc.add_paragraph()
        except Exception:
            pass

    smalls = [views.get(k) for k in ("top", "front", "side")]
    if any(_png_bytes(v) for v in smalls):
        try:
            t = doc.add_table(rows=1, cols=3)
            t.autofit = True
            for i, v in enumerate(smalls):
                png = _png_bytes(v)
                if not png:
                    continue
                cell = t.cell(0, i)
                cell.paragraphs[0].add_run().add_picture(io.BytesIO(png), width=Cm(5.4))
            doc.add_paragraph()
        except Exception:
            pass

def build_docx(data, views=None) -> bytes:
    from docx import Document
    from docx.shared import Pt, Mm
    from docx.enum.table import WD_TABLE_ALIGNMENT

    p = _build_payload(data)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    doc.add_heading("PackAssist — Informe d'embalatge", level=1)
    doc.add_paragraph(p["name"])
    doc.add_paragraph("Document generat automàticament")

    def kv_table(rows, title=None):
        if title:
            doc.add_heading(title, level=2)
        t = doc.add_table(rows=len(rows), cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (k, v) in enumerate(rows):
            t.cell(i, 0).text = str(k)
            t.cell(i, 1).text = str(v)
            for c in (t.cell(i, 0), t.cell(i, 1)):
                for par in c.paragraphs:
                    for run in par.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph()

    # Part information
    part = p["part"]
    kv_table([
        ("Nº de peça", part.get("number") or "—"),
        ("Projecte", part.get("project") or "—"),
        ("Material", part.get("material") or "—"),
        ("Proveïdor", p["supplier"].get("name") or "—"),
        ("Dimensions", p["dims"]),
        ("Pes / peça", _fmt_kg(p["piece_weight"]) if p["piece_weight"] else "—"),
        ("Mètode", p["mode_label"]),
        ("Percentatge d'ocupació", p["fill_rate"]),
    ], "Informació de la peça")

    # 3D renders (same images as the PDF: isometric hero + top/front/side)
    _add_views_docx(doc, views)

    # Pack metrics
    kv_table([
        ("Peces per caixa", _fmt(p["piece_count"])),
        ("Percentatge d'ocupació", p["fill_rate"]),
        ("Caixa (mm)", p["box_dims"]),
        ("Pes total", _fmt_kg(p["total_weight"]) if p["total_weight"] else "—"),
    ], "Dades d'embalatge")

    # Warnings
    if p["interlock_count"] > 0 and p["show_interlock"]:
        warn = doc.add_paragraph()
        run = warn.add_run(f"⚠ {p['interlock_count']} peces no es poden extreure verticalment (entrellaçades)")
        run.bold = True
    if p["tray_count"] > 1:
        doc.add_paragraph(f"Multi-caixa: {p['tray_count']} caixes ({' + '.join(str(x) for x in p['tray_pieces'])}) = {_fmt(p['piece_count'])} peces")

    # Box info
    kv_table([
        ("Dimensions externes (mm)", p["box_dims"]),
        ("Peces / caixa", _fmt(p["piece_count"])),
        ("Percentatge d'ocupació", p["fill_rate"]),
        ("Mètode", p["mode_label"]),
        ("Pes total (kg)", _fmt(p["total_weight"], 3) if p["total_weight"] else "—"),
        ("Pes màxim (kg)", _fmt(p["max_weight"], 3) if p["max_weight"] else "—"),
    ], "Unitat d'embalatge (caixa)")

    # Pallet
    if p["has_pallet"]:
        pal = p["pallet"]
        kv_table([
            ("Dimensions externes (mm)",
             f"{_fmt_mm(pal.get('l'))} × {_fmt_mm(pal.get('w'))} × {_fmt_mm(pal.get('h'))}"),
            ("Pes (kg)", _fmt(pal.get("weight"), 1) if pal.get("weight") is not None else "—"),
            ("Caixes / palet", _fmt(pal.get("boxes")) if pal.get("boxes") is not None else "—"),
        ], "Unitat d'enviament (palet)")

    # Components
    doc.add_heading("Components d'embalatge", level=2)
    if p["has_items"]:
        t = doc.add_table(rows=1 + len(p["items"]), cols=6)
        t.style = "Table Grid"
        headers = ["Descripció", "Material", "L×W×H (mm)", "Qty", "Preu / unitat", "Cost"]
        for j, h in enumerate(headers):
            t.cell(0, j).text = h
        for i, row in enumerate(p["items"], start=1):
            for j, v in enumerate([row["desc"], row["material"], row["dims"],
                                   row["qty"], row["price"], row["cost"]]):
                t.cell(i, j).text = str(v)
        doc.add_paragraph()
    else:
        doc.add_paragraph("Cap component d'embalatge definit")

    # Cost summary
    if p["has_cost"]:
        kv_table([
            ("Cost caixa", _fmt_euro(p["box_cost"]) if p["box_cost"] is not None else "—"),
            ("Embalatge", _fmt_euro(p["packaging_cost"]) if p["packaging_cost"] is not None else "—"),
            ("Transport", _fmt_euro(p["freight_cost"]) if p["freight_cost"] is not None else "—"),
            ("Cost per peça", _fmt_euro(p["cost_per_part"]) if p["cost_per_part"] is not None else "—"),
        ], "Resum de costos")

    # Supplier
    if p["has_supplier"]:
        sup = p["supplier"]
        kv_table([
            ("Nom", sup.get("name") or "—"),
            ("Adreça", sup.get("address") or "—"),
            ("Contacte", sup.get("contact") or "—"),
            ("Telèfon", sup.get("phone") or "—"),
            ("Email", sup.get("email") or "—"),
            ("Funció", sup.get("function") or "—"),
        ], "Informació del proveïdor")

    # Approvals
    if p["has_approvals"]:
        app = p["approvals"]
        kv_table([
            ("Creat per", app.get("createdBy") or "—"),
            ("Concepte — Funció", app.get("conceptFunction") or "—"),
            ("Concepte — Nom", app.get("conceptName") or "—"),
            ("Concepte — Data", app.get("conceptDate") or "—"),
            ("Final — Funció", app.get("finalFunction") or "—"),
            ("Final — Nom", app.get("finalName") or "—"),
            ("Final — Data", app.get("finalDate") or "—"),
        ], "Aprovacions")

    # Comments
    if p["comments"]:
        doc.add_heading("Observacions", level=2)
        doc.add_paragraph(p["comments"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────── XLSX ───────────────────────────

def _add_views_xlsx(ws, views):
    """Embed the PDF's 3D renders in the Informe sheet: hero + the three
    small views, anchored in the free columns right of the key-values."""
    if not views:
        return
    try:
        from openpyxl.drawing.image import Image as XLImage
    except Exception:
        return
    from openpyxl.utils import get_column_letter

    def add_png(data_url, col, row, px_w):
        png = _png_bytes(data_url)
        if not png:
            return 0
        try:
            from PIL import Image as PILImage
            with PILImage.open(io.BytesIO(png)) as im:
                w, h = im.size
            px_h = max(1, int(px_w * h / w)) if w > 0 else px_w
        except Exception:
            px_h = int(px_w * 0.4)
        img = XLImage(io.BytesIO(png))
        img.width = px_w
        img.height = px_h
        ws.add_image(img, f"{get_column_letter(col)}{row}")
        return max(1, px_h // 16 + 2)

    row = ws.max_row + 2 if ws.max_row > 1 else 2
    row += add_png(views.get("hero"), 4, row, 460)
    for v in (views.get("top"), views.get("front"), views.get("side")):
        row += 1
        row += add_png(v, 4, row, 220)

def build_xlsx(data, views=None) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    p = _build_payload(data)
    wb = Workbook()

    # Sheet 1: informe (key-value)
    ws = wb.active
    ws.title = "Informe"
    bold = Font(bold=True)
    hdr = PatternFill("solid", fgColor="0047AB")
    white = Font(color="FFFFFF", bold=True)

    def kv_sheet(ws, title, rows):
        ws.cell(row=ws.max_row + 1 if ws.max_row > 1 else 1, column=1, value=title).font = Font(bold=True, size=13)
        start = ws.max_row + 1
        for i, (k, v) in enumerate(rows):
            ws.cell(row=start + i, column=1, value=k).font = bold
            ws.cell(row=start + i, column=2, value=v)
        ws.cell(row=start + len(rows) + 1, column=1, value="")
        ws.column_dimensions["A"].width = 30
        ws.column_dimensions["B"].width = 45

    ws.append(["Informe d'embalatge — PackAssist"])
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = p["name"]
    ws.append([])

    kv_sheet(ws, "Informació de la peça", [
        ("Nº de peça", p["part"].get("number") or "—"),
        ("Projecte", p["part"].get("project") or "—"),
        ("Material", p["part"].get("material") or "—"),
        ("Proveïdor", p["supplier"].get("name") or "—"),
        ("Dimensions", p["dims"]),
        ("Pes / peça", _fmt_kg(p["piece_weight"]) if p["piece_weight"] else "—"),
        ("Mètode", p["mode_label"]),
        ("Percentatge d'ocupació", p["fill_rate"]),
    ])
    kv_sheet(ws, "Dades d'embalatge", [
        ("Peces per caixa", _fmt(p["piece_count"])),
        ("Percentatge d'ocupació", p["fill_rate"]),
        ("Caixa (mm)", p["box_dims"]),
        ("Pes total", _fmt_kg(p["total_weight"]) if p["total_weight"] else "—"),
        ("Pes màxim (kg)", _fmt(p["max_weight"], 3) if p["max_weight"] else "—"),
    ])
    kv_sheet(ws, "Unitat d'embalatge (caixa)", [
        ("Dimensions externes (mm)", p["box_dims"]),
        ("Peces / caixa", _fmt(p["piece_count"])),
        ("Percentatge d'ocupació", p["fill_rate"]),
        ("Mètode", p["mode_label"]),
        ("Pes total (kg)", _fmt(p["total_weight"], 3) if p["total_weight"] else "—"),
    ])

    if p["interlock_count"] > 0 and p["show_interlock"]:
        warn_row = ws.max_row + 1
        ws.cell(row=warn_row, column=1,
                value=f"AVÍS: {p['interlock_count']} peces no es poden extreure verticalment (entrellaçades)")
        ws.cell(row=warn_row, column=1).font = Font(bold=True, color="B91C1C")
        ws.append([])
    if p["tray_count"] > 1:
        ws.append([f"Multi-caixa: {p['tray_count']} caixes "
                   f"({' + '.join(str(x) for x in p['tray_pieces'])}) = {_fmt(p['piece_count'])} peces"])
        ws.append([])

    # 3D renders (same images as the PDF), anchored below the key-values.
    _add_views_xlsx(ws, views)

    # Sheet 2: components table
    ws2 = wb.create_sheet("Components")
    headers = ["Descripció", "Material", "L", "W", "H", "Qty", "Preu / unitat", "Cost"]
    for j, h in enumerate(headers, start=1):
        c = ws2.cell(row=1, column=j, value=h)
        c.font = white
        c.fill = hdr
    raw_items = data.get("items") or []
    for i, it in enumerate(raw_items, start=2):
        qty = it.get("qty")
        price = it.get("price")
        cost_v = (qty or 0) * (price or 0) if qty is not None and price is not None else None
        ws2.cell(row=i, column=1, value=it.get("desc") or "")
        ws2.cell(row=i, column=2, value=it.get("material") or "")
        ws2.cell(row=i, column=3, value=it.get("l"))
        ws2.cell(row=i, column=4, value=it.get("w"))
        ws2.cell(row=i, column=5, value=it.get("h"))
        ws2.cell(row=i, column=6, value=qty)
        ws2.cell(row=i, column=7, value=price)
        ws2.cell(row=i, column=8, value=cost_v)
    for col, w in zip("ABCDEFGH", (45, 18, 9, 9, 9, 9, 14, 14)):
        ws2.column_dimensions[col].width = w

    # Sheet 3: cost + supplier + approvals + comments
    ws3 = wb.create_sheet("Costos i proveïdor")
    p3 = _build_payload(data)
    kv_sheet(ws3, "Resum de costos", [
        ("Cost caixa", _fmt_euro(p3["box_cost"]) if p3["box_cost"] is not None else "—"),
        ("Embalatge", _fmt_euro(p3["packaging_cost"]) if p3["packaging_cost"] is not None else "—"),
        ("Transport", _fmt_euro(p3["freight_cost"]) if p3["freight_cost"] is not None else "—"),
        ("Cost per peça", _fmt_euro(p3["cost_per_part"]) if p3["cost_per_part"] is not None else "—"),
    ])
    if p3["has_supplier"]:
        sup = p3["supplier"]
        kv_sheet(ws3, "Informació del proveïdor", [
            ("Nom", sup.get("name") or "—"),
            ("Adreça", sup.get("address") or "—"),
            ("Contacte", sup.get("contact") or "—"),
            ("Telèfon", sup.get("phone") or "—"),
            ("Email", sup.get("email") or "—"),
            ("Funció", sup.get("function") or "—"),
        ])
    if p3["has_approvals"]:
        app = p3["approvals"]
        kv_sheet(ws3, "Aprovacions", [
            ("Creat per", app.get("createdBy") or "—"),
            ("Concepte — Funció", app.get("conceptFunction") or "—"),
            ("Concepte — Nom", app.get("conceptName") or "—"),
            ("Concepte — Data", app.get("conceptDate") or "—"),
            ("Final — Funció", app.get("finalFunction") or "—"),
            ("Final — Nom", app.get("finalName") or "—"),
            ("Final — Data", app.get("finalDate") or "—"),
        ])
    if p3["comments"]:
        ws3.append([])
        ws3.append(["Observacions"])
        ws3.cell(row=ws3.max_row, column=1).font = bold
        ws3.append([p3["comments"]])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()